from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
UNORDERED_RE = re.compile(r"^(\s*)-\s+(.*)$")


def _clean_plain_inline(text: str) -> str:
    """Strip lightweight Markdown markers for plain-text Word output."""
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text.strip()


def _clean_inline(text: str) -> str:
    """Preserve inline markdown markers that are rendered later."""
    return text.strip()


def _append_inline_runs(paragraph, text: str) -> None:
    """Render lightweight inline markdown for bold and code spans."""
    i = 0
    buffer = []
    bold = False
    code = False

    def flush_buffer() -> None:
        nonlocal buffer
        if not buffer:
            return
        run = paragraph.add_run("".join(buffer))
        run.bold = bold
        if code:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        buffer = []

    while i < len(text):
        if text.startswith("**", i):
            flush_buffer()
            bold = not bold
            i += 2
            continue
        if text[i] == "`":
            flush_buffer()
            code = not code
            i += 1
            continue
        buffer.append(text[i])
        i += 1

    flush_buffer()


def _add_paragraph_with_inline(document: Document, text: str, style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    _append_inline_runs(paragraph, text.strip())
    return paragraph


def _add_heading_with_inline(document: Document, text: str, level: int):
    paragraph = document.add_heading(level=level)
    _append_inline_runs(paragraph, text.strip())
    return paragraph


def _write_code_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text.rstrip("\n"))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def _write_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    table = document.add_table(rows=len(normalized), cols=width)
    table.style = "Table Grid"
    for i, row in enumerate(normalized):
        for j, cell_text in enumerate(row):
            table.rows[i].cells[j].text = _clean_plain_inline(cell_text)


def _parse_table_line(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def _is_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|").replace("-", "").replace(":", "").replace(" ", "")
    return stripped == ""


def _add_markdown_to_document(document: Document, markdown_text: str) -> None:
    lines = markdown_text.splitlines()
    in_code_block = False
    table_buffer: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            _write_table(document, table_buffer)
            table_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_table()
            in_code_block = not in_code_block
            continue

        if in_code_block:
            _write_code_paragraph(document, line)
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            parsed = _parse_table_line(stripped)
            if _is_table_separator(stripped):
                continue
            table_buffer.append(parsed)
            continue

        flush_table()

        if not stripped:
            document.add_paragraph("")
            continue

        heading_match = HEADING_RE.match(line)
        if heading_match:
            hashes, text = heading_match.groups()
            level = min(len(hashes), 6)
            _add_heading_with_inline(document, _clean_inline(text), level=level)
            continue

        ordered_match = ORDERED_RE.match(line)
        if ordered_match:
            spaces, _, text = ordered_match.groups()
            level = min(len(spaces) // 2 + 1, 3)
            style = "List Number" if level == 1 else f"List Number {level}"
            _add_paragraph_with_inline(document, _clean_inline(text), style=style)
            continue

        unordered_match = UNORDERED_RE.match(line)
        if unordered_match:
            spaces, text = unordered_match.groups()
            level = min(len(spaces) // 2 + 1, 3)
            style = "List Bullet" if level == 1 else f"List Bullet {level}"
            _add_paragraph_with_inline(document, _clean_inline(text), style=style)
            continue

        if stripped.startswith(">"):
            _add_paragraph_with_inline(
                document,
                _clean_inline(stripped.lstrip(">")),
                style="Intense Quote",
            )
            continue

        _add_paragraph_with_inline(document, _clean_inline(stripped))

    flush_table()


def markdown_file_to_docx(source: Path, destination: Path) -> Path:
    document = Document()
    _add_markdown_to_document(document, source.read_text(encoding="utf-8"))
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)
    return destination


def consolidate_markdown_to_docx(sources: Iterable[Path], destination: Path, title: str) -> Path:
    document = Document()
    _add_heading_with_inline(document, title, level=0)

    source_list = list(sources)
    for index, source in enumerate(source_list):
        _add_heading_with_inline(document, source.stem, level=1)
        _add_markdown_to_document(document, source.read_text(encoding="utf-8"))
        if index < len(source_list) - 1:
            paragraph = document.add_paragraph()
            paragraph.add_run().add_break(WD_BREAK.PAGE)

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)
    return destination


def _run_default_batch() -> None:
    docs_dir = Path("docs/02_architecture")
    overview = docs_dir / "CURRENT_CODEBASE_OVERVIEW_2026-03-12_215446.md"
    theory = docs_dir / "CURRENT_TECHNICAL_THEORY_OF_OPERATION_2026-03-12_215446.md"

    overview_outputs = [
        docs_dir / "CURRENT_CODEBASE_OVERVIEW_2026-03-12_215446.docx",
        docs_dir / "HybridRAG3_Codebase_Overview.docx",
    ]
    theory_outputs = [
        docs_dir / "CURRENT_TECHNICAL_THEORY_OF_OPERATION_2026-03-12_215446.docx",
        docs_dir / "HybridRAG3_Technical_Theory_of_Operation.docx",
    ]
    synthesis_outputs = [
        docs_dir / "CURRENT_CODEBASE_SYNTHESIS_2026-03-12_215446.docx",
        docs_dir / "HybridRAG3_Codebase_Synthesis.docx",
    ]

    for destination in overview_outputs:
        markdown_file_to_docx(overview, destination)

    for destination in theory_outputs:
        markdown_file_to_docx(theory, destination)

    for destination in synthesis_outputs:
        consolidate_markdown_to_docx(
            [overview, theory],
            destination,
            title="HybridRAG3 Current Codebase Synthesis",
        )


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert markdown docs to docx.")
    parser.add_argument("sources", nargs="*", help="Markdown source files.")
    parser.add_argument("--output", help="Destination .docx path.")
    parser.add_argument("--title", help="Title for consolidated output.")
    args = parser.parse_args()

    if args.sources and args.output:
        source_paths = [Path(path) for path in args.sources]
        destination = Path(args.output)
        if len(source_paths) == 1:
            markdown_file_to_docx(source_paths[0], destination)
        else:
            title = args.title or destination.stem.replace("_", " ")
            consolidate_markdown_to_docx(source_paths, destination, title)
        return

    _run_default_batch()


if __name__ == "__main__":
    main()
