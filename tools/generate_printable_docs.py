"""
Generate print-ready Word (.docx) documents from key markdown files.
FILE: tools/generate_printable_docs.py

Run:   python tools/generate_printable_docs.py
Output: docs/_printable/*.docx  (gitignored, generated locally)

The .md source stays in git for transfer. Run this script on any machine
to produce formatted Word docs for printing or studying offline.
"""
import os
import re
import sys

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ---------------------------------------------------------------------------
# CONFIGURATION
# ---------------------------------------------------------------------------

OUTPUT_DIR = os.path.join(PROJECT_ROOT, "docs", "_printable")

# (source_path_relative_to_project, output_filename, category_color_hex)
DOCS_TO_CONVERT = [
    # Architecture (blue)
    ("docs/02_architecture/THEORY_OF_OPERATION_RevA.md",
     "01_Theory_of_Operation.docx", "1F4E79"),
    ("docs/02_architecture/TECHNICAL_THEORY_OF_OPERATION_RevC.md",
     "02_Technical_Theory_of_Operation.docx", "1F4E79"),
    ("docs/02_architecture/THEORY_OF_OPERATION_EXECUTIVE.md",
     "03_Executive_Summary.docx", "1F4E79"),
    ("docs/02_architecture/SECURITY_THEORY_OF_OPERATION_RevA.md",
     "04_Security_Theory_of_Operation.docx", "8B0000"),
    ("docs/02_architecture/CAPABILITIES.md",
     "05_Capabilities.docx", "1F4E79"),
    ("docs/02_architecture/PIPELINE_FLOW_AND_BOTTLENECKS.md",
     "06_Pipeline_Flow_and_Bottlenecks.docx", "1F4E79"),
    ("docs/02_architecture/SOFTWARE_TRADEOFFS.md",
     "07_Software_Tradeoffs.docx", "1F4E79"),
    ("docs/02_architecture/COMPETITIVE_DIFFERENTIATION.md",
     "08_Competitive_Differentiation.docx", "1F4E79"),
    ("docs/02_architecture/CODEBASE_SIZE_BREAKDOWN.md",
     "09_Codebase_Size_Breakdown.docx", "1F4E79"),
    ("docs/02_architecture/PRODUCTION_SCALE_ESTIMATE.md",
     "10_Production_Scale_Estimate.docx", "1F4E79"),
    ("docs/02_architecture/ROI_AND_BENEFITS.md",
     "11_ROI_and_Benefits.docx", "1F4E79"),
    ("docs/02_architecture/ELEVATOR_SPEECH.md",
     "12_Elevator_Speech.docx", "1F4E79"),
    ("docs/02_architecture/FORMAT_SUPPORT.md",
     "13_Format_Support.docx", "1F4E79"),
    ("docs/02_architecture/INTERFACES.md",
     "14_Interfaces.docx", "1F4E79"),
    ("docs/02_architecture/SYSTEM_UPGRADE_ROADMAP.md",
     "15_System_Upgrade_Roadmap.docx", "1F4E79"),
    ("docs/02_architecture/SOFTWARE_HISTORY_AND_SCALABILITY_PLAN_2026-03-02.md",
     "16_Software_History_and_Scalability.docx", "1F4E79"),
    ("docs/02_architecture/ARCHITECTURE_DIAGRAM.md",
     "17_Architecture_Diagram.docx", "1F4E79"),
    ("docs/02_architecture/TWO_REPO_STRATEGY.md",
     "18_Two_Repo_Strategy.docx", "1F4E79"),
    # Guides (green)
    ("docs/03_guides/USER_GUIDE.md",
     "20_User_Guide.docx", "1B5E20"),
    ("docs/03_guides/GUI_GUIDE.md",
     "21_GUI_Guide.docx", "1B5E20"),
    ("docs/03_guides/GLOSSARY.md",
     "22_Glossary.docx", "1B5E20"),
    ("docs/03_guides/SHORTCUT_SHEET.md",
     "23_Shortcut_Sheet.docx", "1B5E20"),
    # Setup (orange)
    ("docs/01_setup/INSTALL_AND_SETUP.md",
     "30_Install_and_Setup.docx", "BF6900"),
    ("docs/01_setup/MANUAL_INSTALL.md",
     "31_Manual_Install.docx", "BF6900"),
    ("docs/01_setup/WORK_LAPTOP_VENV_SETUP.md",
     "32_Work_Laptop_Setup.docx", "BF6900"),
    # Security (dark red)
    ("docs/05_security/waiver_reference_sheet.md",
     "40_Waiver_Reference_Sheet.docx", "8B0000"),
    ("docs/05_security/DEFENSE_MODEL_AUDIT.md",
     "41_Model_Audit.docx", "8B0000"),
    ("docs/05_security/GIT_REPO_RULES.md",
     "42_Git_Repo_Rules.docx", "8B0000"),
    # Inventory (purple)
    ("docs/Development_Inventory.md",
     "50_Development_Inventory.docx", "4A148C"),
    # Demo (teal)
    ("docs/04_demo/DEMO_GUIDE.md",
     "60_Demo_Guide.docx", "006064"),
    ("docs/04_demo/DEMO_PREP.md",
     "61_Demo_Prep.docx", "006064"),
    # Learning (indigo)
    ("docs/08_learning/STUDY_GUIDE.md",
     "70_Study_Guide.docx", "1A237E"),
    ("docs/08_learning/RAG_LANDSCAPE_2026.md",
     "71_RAG_Landscape_2026.docx", "1A237E"),
]


# ---------------------------------------------------------------------------
# INLINE FORMATTING PARSER
# ---------------------------------------------------------------------------

# Regex to split on bold/italic/code/link markers
_INLINE_RE = re.compile(
    r'(\*\*\*(.+?)\*\*\*'       # ***bold italic***
    r'|\*\*(.+?)\*\*'            # **bold**
    r'|\*(.+?)\*'                # *italic*
    r'|`([^`]+)`'                # `code`
    r'|\[([^\]]+)\]\([^)]+\))'   # [link](url) -> just text
)


def _add_formatted_runs(paragraph, text, base_size=10):
    """Parse inline markdown and add formatted runs to a paragraph."""
    if not text:
        return
    pos = 0
    for m in _INLINE_RE.finditer(text):
        # Add plain text before this match
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.size = Pt(base_size)
        # Determine which group matched
        if m.group(2):  # bold italic
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = True
            run.font.size = Pt(base_size)
        elif m.group(3):  # bold
            run = paragraph.add_run(m.group(3))
            run.bold = True
            run.font.size = Pt(base_size)
        elif m.group(4):  # italic
            run = paragraph.add_run(m.group(4))
            run.italic = True
            run.font.size = Pt(base_size)
        elif m.group(5):  # inline code
            run = paragraph.add_run(m.group(5))
            run.font.name = "Consolas"
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        elif m.group(6):  # link text
            run = paragraph.add_run(m.group(6))
            run.font.size = Pt(base_size)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        pos = m.end()
    # Remaining plain text
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = Pt(base_size)


# ---------------------------------------------------------------------------
# TABLE HELPERS
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color):
    """Apply background color to a table cell."""
    shading = parse_xml(
        '<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), hex_color)
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_text(cell, text, bold=False, size=9, color=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    _add_formatted_runs(p, text, base_size=size)
    if bold:
        for run in p.runs:
            run.bold = True
    if color:
        for run in p.runs:
            run.font.color.rgb = color


def _parse_table_row(line):
    """Parse a markdown table row into cells."""
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def _is_separator_row(line):
    """Check if a line is a table separator (|---|---|)."""
    return bool(re.match(r'^[\s|:-]+$', line.replace("-", "")))


# ---------------------------------------------------------------------------
# MARKDOWN TO DOCX CONVERTER
# ---------------------------------------------------------------------------

def convert_md_to_docx(md_path, docx_path, accent_hex):
    """Convert a markdown file to a formatted Word document."""
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # Page setup: letter size, reasonable margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    accent_rgb = RGBColor(
        int(accent_hex[0:2], 16),
        int(accent_hex[2:4], 16),
        int(accent_hex[4:6], 16),
    )

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # Title style
    for level in range(1, 5):
        hstyle = doc.styles["Heading %d" % level]
        hstyle.font.color.rgb = accent_rgb if level <= 2 else RGBColor(0x33, 0x33, 0x33)

    i = 0
    in_code_block = False
    code_lines = []
    table_rows = []
    table_started = False
    list_buffer = []
    list_type = None  # 'bullet' or 'number'

    def _flush_list():
        """Write buffered list items to the document."""
        nonlocal list_buffer, list_type
        if not list_buffer:
            return
        for item_text in list_buffer:
            style_name = "List Bullet" if list_type == "bullet" else "List Number"
            p = doc.add_paragraph(style=style_name)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            _add_formatted_runs(p, item_text, base_size=10)
        list_buffer = []
        list_type = None

    def _flush_table():
        """Write buffered table rows to the document."""
        nonlocal table_rows, table_started
        if not table_rows:
            return
        # First row is header
        num_cols = len(table_rows[0])
        table = doc.add_table(rows=0, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"

        for row_idx, cells in enumerate(table_rows):
            row = table.add_row()
            # Pad or truncate cells to match column count
            while len(cells) < num_cols:
                cells.append("")
            for col_idx, cell_text in enumerate(cells[:num_cols]):
                cell = row.cells[col_idx]
                if row_idx == 0:
                    _shade_cell(cell, accent_hex)
                    _set_cell_text(cell, cell_text, bold=True, size=9,
                                   color=RGBColor(0xFF, 0xFF, 0xFF))
                else:
                    if row_idx % 2 == 0:
                        _shade_cell(cell, "F2F2F2")
                    _set_cell_text(cell, cell_text, size=9)

        doc.add_paragraph()  # spacing after table
        table_rows = []
        table_started = False

    while i < len(lines):
        line = lines[i].rstrip("\n").rstrip("\r")

        # --- Code block ---
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block
                _flush_list()
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                # Light gray background via shading
                shading = parse_xml(
                    '<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w'))
                )
                p._element.get_or_add_pPr().append(shading)
                code_lines = []
                in_code_block = False
            else:
                _flush_list()
                _flush_table()
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # --- HTML comments / hidden tags ---
        if line.strip().startswith("<!--"):
            # Skip comment lines
            while i < len(lines) and "-->" not in lines[i]:
                i += 1
            i += 1
            continue

        # --- Blank line ---
        if not line.strip():
            _flush_list()
            if table_started:
                _flush_table()
            i += 1
            continue

        # --- Horizontal rule ---
        if re.match(r'^-{3,}$', line.strip()) or re.match(r'^\*{3,}$', line.strip()):
            _flush_list()
            _flush_table()
            # Add a subtle separator
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("_" * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue

        # --- Headings ---
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            _flush_list()
            _flush_table()
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Remove trailing # characters
            text = re.sub(r'\s*#+\s*$', '', text)
            heading_level = min(level, 4)
            h = doc.add_heading(level=heading_level)
            _add_formatted_runs(h, text, base_size=max(16 - level * 2, 10))
            i += 1
            continue

        # --- Table row ---
        if "|" in line and not line.strip().startswith(">"):
            cells = _parse_table_row(line)
            if _is_separator_row(line):
                # Skip separator row (|---|---|)
                i += 1
                continue
            _flush_list()
            if not table_started:
                table_started = True
                table_rows = []
            table_rows.append(cells)
            i += 1
            continue

        # --- Bullet list ---
        bullet_match = re.match(r'^(\s*)[*\-+]\s+(.*)', line)
        if bullet_match:
            _flush_table()
            if list_type != "bullet":
                _flush_list()
                list_type = "bullet"
            list_buffer.append(bullet_match.group(2))
            i += 1
            continue

        # --- Numbered list ---
        num_match = re.match(r'^(\s*)\d+[.)]\s+(.*)', line)
        if num_match:
            _flush_table()
            if list_type != "number":
                _flush_list()
                list_type = "number"
            list_buffer.append(num_match.group(2))
            i += 1
            continue

        # --- Blockquote ---
        if line.strip().startswith(">"):
            _flush_list()
            _flush_table()
            text = re.sub(r'^>\s*', '', line.strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # --- Regular paragraph ---
        _flush_list()
        _flush_table()
        p = doc.add_paragraph()
        _add_formatted_runs(p, line, base_size=10)
        i += 1

    # Flush anything remaining
    _flush_list()
    _flush_table()
    if code_lines:
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = Pt(8)

    # Add footer with page numbers
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("HybridRAG v3")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        # Page number field
        fp.add_run("  |  Page ")
        fld_xml = (
            '<w:fldSimple {} w:instr=" PAGE "><w:r><w:t>0</w:t></w:r>'
            '</w:fldSimple>'
        ).format(nsdecls('w'))
        fp._element.append(parse_xml(fld_xml))

    doc.save(docx_path)


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print("Generating hardcopy Word documents...")
    print("Output: %s" % OUTPUT_DIR)
    print()

    converted = 0
    skipped = 0

    for rel_path, out_name, color in DOCS_TO_CONVERT:
        src = os.path.join(PROJECT_ROOT, rel_path)
        dst = os.path.join(OUTPUT_DIR, out_name)
        if not os.path.exists(src):
            print("  [SKIP] %s (not found)" % rel_path)
            skipped += 1
            continue
        try:
            convert_md_to_docx(src, dst, color)
            print("  [OK]   %s" % out_name)
            converted += 1
        except Exception as e:
            print("  [FAIL] %s -- %s" % (out_name, e))
            skipped += 1

    print()
    print("[OK] Generated %d documents (%d skipped)" % (converted, skipped))
    print("     Location: %s" % OUTPUT_DIR)
    print()
    print("To print: open the folder in Explorer and print what you need.")
    print("These files are gitignored -- regenerate on any machine with:")
    print("  python tools/generate_printable_docs.py")


if __name__ == "__main__":
    main()
